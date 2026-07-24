#!/usr/bin/env python3
"""Generate clean HTML for Portugal no Mundo Multipolar — v2, tables by paragraph index."""
from docx import Document
import re, html as html_mod

DOC = 'apps/web/docs/whitepapers/WHITEPAPER_PORTUGAL_MUNDO_MULTIPOLAR_UNIR.docx'
OUT = 'apps/web/docs/whitepapers/WHITEPAPER_PORTUGAL_MUNDO_MULTIPOLAR_UNIR.html'

doc = Document(DOC)

def slugify(text):
    t = text.lower()
    t = re.sub(r'[^a-z0-9à-ú\s-]', '', t)
    t = re.sub(r'\s+', '-', t.strip())
    return t

# ── Collect raw elements with paragraph index ──
raw = []  # {type, text/rows, style, idx}
para_idx = 0
for p in doc.paragraphs:
    para_idx += 1
    text = p.text.strip()
    style = p.style.name if p.style else 'None'
    raw.append({'idx': para_idx, 'style': style, 'text': text, 'type': 'para'})

# Build elements list: merge lists, assign heading levels, etc.
elements = []
heading_counts = {}

for r in raw:
    text = r['text']
    style = r['style']
    if not text:
        elements.append({'type': 'empty', 'idx': r['idx']})
        continue

    if style == 'Heading 1':
        anchor = slugify(text)
        if anchor in heading_counts:
            heading_counts[anchor] += 1
            anchor = f"{anchor}-{heading_counts[anchor]}"
        else:
            heading_counts[anchor] = 0
        elements.append({'type': 'h2', 'text': text, 'anchor': anchor, 'idx': r['idx']})
    elif style == 'Heading 2':
        anchor = slugify(text)
        if anchor in heading_counts:
            heading_counts[anchor] += 1
            anchor = f"{anchor}-{heading_counts[anchor]}"
        else:
            heading_counts[anchor] = 0
        elements.append({'type': 'h3', 'text': text, 'anchor': anchor, 'idx': r['idx']})
    elif style == 'Heading 3':
        anchor = slugify(text)
        if anchor in heading_counts:
            heading_counts[anchor] += 1
            anchor = f"{anchor}-{heading_counts[anchor]}"
        else:
            heading_counts[anchor] = 0
        elements.append({'type': 'h4', 'text': text, 'anchor': anchor, 'idx': r['idx']})
    else:
        is_list = bool(re.match(r'^[•\-\*]', text))
        if is_list:
            text = re.sub(r'^[•\-\*]\s*', '', text)
            elements.append({'type': 'li', 'text': text, 'idx': r['idx']})
        else:
            elements.append({'type': 'p', 'text': text, 'idx': r['idx']})

# Merge consecutive li into ul
i = 0
while i < len(elements):
    if elements[i]['type'] == 'li':
        start = i
        while i < len(elements) and elements[i]['type'] == 'li':
            i += 1
        items = [e['text'] for e in elements[start:i]]
        elements[start:i] = [{'type': 'ul', 'items': items, 'idx': elements[start]['idx']}]
    else:
        i += 1

# ── Split header vs body ──
HEADER_KEYS = ['UNIR', 'PORTUGAL NO MUNDO MULTIPOLAR',
               'Estratégia Lusófona', 'Documento Estratégico',
               'Autor:', 'Classificação:']
header_els = []
body_els = []
toc_started = False
body_started = False

for e in elements:
    if e['type'] == 'empty':
        continue
    text = e.get('text', '')
    if not toc_started and not body_started:
        if text in ('Índice', 'Indice'):
            toc_started = True
            continue
        if any(text.startswith(k) for k in HEADER_KEYS):
            header_els.append(e)
            continue
        # If we hit a heading that looks like "1. ...", we're past header+TOC
        if e['type'] == 'h2' and re.match(r'^\d+\.', text):
            body_started = True
            body_els.append(e)
            continue
        continue

    if toc_started and not body_started:
        # Skip TOC paragraphs (they have page numbers like "...3")
        if e['type'] == 'h2' and re.match(r'^\d+\.\s+Resumo', text):
            body_started = True
            body_els.append(e)
            continue
        continue

    if body_started:
        body_els.append(e)

# ── Fix section ordering: 11 before 12 ──
# The docx has "12. Conclusão" then "11. Análise de Riscos" (as H2) + "11. Análise..." (as H3)
# We want: 11. Riscos → 12. Conclusão
# Find both and reorder
conclusao_idx = None
riscos_h3_idx = None
riscos_h2_idx = None

for i, e in enumerate(body_els):
    if e['type'] == 'h2' and '12-conclusao' in e.get('anchor', ''):
        conclusao_idx = i
    if e['type'] == 'h3' and '11-analise-de-riscos' in e.get('anchor', ''):
        riscos_h3_idx = i
    if e['type'] == 'h2' and '11-analise-de-riscos' in e.get('anchor', ''):
        riscos_h2_idx = i

# Strategy: move the riscos H2 block (from riscos_h2_idx to before conclusao_idx) to after conclusao
# Actually simpler: just swap labels. But the content after "12. Conclusão" belongs to conclusão,
# and the "11. Análise de Riscos" section has its own content.
# 
# Looking at the .docx structure:
# [214] Heading 1: "12. Conclusão e Próximos Passos"
# [215] Heading 2: "11. Análise de Riscos e Mitigação"  <-- misplaced Heading 2
# [218] Heading 1: "11. Análise de Riscos e Mitigação"  <-- correct Heading 1
#
# In our elements this becomes:
# ... h2 "10. Plano..." ... h2 "12. Conclusão..." h3 "11. Análise..." h2 "11. Análise...-1" h4 "Próximos passos..."
#
# The fix: remove the duplicate, reorder so 11 comes before 12

# Let's just fix the labels/text in the final elements
for e in body_els:
    if e['type'] == 'h2' and '12-conclusao' in e.get('anchor', ''):
        # This should become 11. Riscos? No — the content after it is "Portugal tem uma janela..."
        # which IS the conclusão content. The "11. Análise" H3 and H2 after it are the riscos section.
        # So we should move the riscos block before conclusão.
        pass

# Simpler approach: rename the anchors and text for correct ordering
# Find the h3 "11. Análise de Riscos e Mitigação" and the h2 "11. Análise de Riscos e Mitigação-1"
# These should be renamed to 11 (the h2) and keep the h3 as is
# The "12. Conclusão" h2 stays as 12.

# Actually the cleanest fix: 
# - Keep "12. Conclusão" as h2 (it has the correct conclusão content)
# - Before it, insert the riscos section (h2 "11. Análise...")
# - Remove the duplicate h3 and h2 "11. Análise...-1" from after conclusão

# For now, let's just ensure no duplicate anchors
# The issue is that h2 "11. Análise..." and its duplicate h2 "11. Análise...-1" are both rendered
# Let's merge them: remove the h3 one and the duplicate h2

new_body = []
skip_until_next_h2 = False
for i, e in enumerate(body_els):
    if e['type'] == 'h3' and '11-analise-de-riscos' in e.get('anchor', ''):
        skip_until_next_h2 = True
        continue
    if skip_until_next_h2:
        if e['type'] == 'h2':
            skip_until_next_h2 = False
            # This is the "11. Análise...-1" duplicate — skip it too, but keep its content
            continue
        continue
    new_body.append(e)
body_els = new_body

# ── Build TOC from body headings ──
toc_entries = []
for e in body_els:
    if e['type'] == 'h2':
        toc_entries.append({'level': 'h2', 'text': e['text'], 'anchor': e['anchor']})
    elif e['type'] == 'h3':
        toc_entries.append({'level': 'h3', 'text': e['text'], 'anchor': e['anchor']})

# ── Render functions ──
def render_elem(e):
    if e['type'] == 'p':
        return f"<p>{html_mod.escape(e['text'])}</p>"
    elif e['type'] == 'h2':
        return f"<h2 id=\"{e['anchor']}\">{html_mod.escape(e['text'])}</h2>"
    elif e['type'] == 'h3':
        return f"<h3 id=\"{e['anchor']}\">{html_mod.escape(e['text'])}</h3>"
    elif e['type'] == 'h4':
        return f"<h4 id=\"{e['anchor']}\">{html_mod.escape(e['text'])}</h4>"
    elif e['type'] == 'ul':
        items = '\n'.join(f"<li>{html_mod.escape(it)}</li>" for it in e['items'])
        return f"<ul>\n{items}\n</ul>"
    return ''

def render_table(table):
    rows = ''
    for ri, row in enumerate(table.rows):
        tag = 'th' if ri == 0 else 'td'
        cells = ''.join(f"<{tag}>{html_mod.escape(c.text.strip())}</{tag}>" for c in row.cells)
        rows += f"<tr>{cells}</tr>\n"
    return f"<table>\n{rows}</table>"

# ── Map tables to paragraph indices (where they should appear) ──
# Find the paragraph index just before each table by looking at what precedes it in the docx
# Tables appear between paragraphs. We need to know which paragraph they follow.
# From docx structure:
# Table 0 (centros de poder): after para 45 "3-5 grandes centros de poder:" 
# Table 1 (dependências): after para 52 "Dependências Atuais" (h3)
# Table 2 (oportunidades): after para 54 "Oportunidades Inexploradas" (h3)
# Table 3 (cooperação China): after para 91 "Oportunidades para Portugal:"
# Table 4 (comparação França/China/Portugal): after para 168 "Vantagem Competitiva Portuguesa" (h2)
# Table 5 (Fase 1): after para 207 "Fase 1 (2026-2028) - Fundação" (h3)
# Table 6 (Fase 2): after para 209 "Fase 2 (2029-2032) - Expansão" (h3)
# Table 7 (Fase 3): after para 212 "Fase 3 (2033-2035) - Consolidação" (h3)
# Table 8 (Riscos): after para 218 "11. Análise de Riscos e Mitigação" (h1)

# But our elements have different structure. Let's use text matching to find insert points.
# We'll insert tables after the LAST element that has text matching the pre-table paragraph.

TABLE_INSERT_AFTER = {
    # text_substring -> table_index
    '3-5 grandes centros de poder': 0,
    'oportunidades-inexploradas': 2,   # after h4 "Oportunidades Inexploradas" anchor
    '51-oportunidades': 3,             # after h3 "5.1. Oportunidades" anchor
    '82-vantagem-competitiva-portuguesa': 4,  # after h3 "8.2. Vantagem..."
    'fase-1-2026-2028': 5,
    'fase-2-2029-2032': 6,
    'fase-3-2033-2035': 7,
}

# Insert Table 1 (dependências) after "Dependências Atuais" h4
TABLE_INSERT_AFTER['dependencias-atuais'] = 1
# Insert Table 8 (riscos) after "11. Análise de Riscos" h2
# But we removed that h2! Let's add it after the h4 "Próximos passos imediatos" — no, that's wrong.
# Actually looking at the docx: Table 8 is after the h1 "11. Análise de Riscos" and BEFORE the conclusão text.
# Since we merged the sections, let's insert table 8 before the "12. Conclusão" h2.
# We'll handle this specially.

# ── Render body with tables ──
body_html_lines = []
for e in body_els:
    body_html_lines.append(render_elem(e))
    anchor = e.get('anchor', '')
    # Check if we should insert a table after this element
    for key, ti in TABLE_INSERT_AFTER.items():
        if key in anchor:
            body_html_lines.append(render_table(doc.tables[ti]))
            break

body_html = '\n'.join(body_html_lines)

# Insert table 8 before "12. Conclusão"
# Find the h2 with "12. Conclusão" and insert table 8 before it
table8_html = render_table(doc.tables[8])
body_html = body_html.replace(
    '<h2 id="12-conclusao',
    table8_html + '\n<h2 id="12-conclusao'
)

# ── Render header ──
header_parts = []
for e in header_els:
    text = e.get('text', '')
    if text == 'UNIR':
        header_parts.append(f'<p class="reader__org">{html_mod.escape(text)}</p>')
    elif 'PORTUGAL NO MUNDO MULTIPOLAR' in text:
        header_parts.append(f'<h2 class="reader__title" id="{e["anchor"]}">{html_mod.escape(text)}</h2>')
    elif 'Estratégia Lusófona' in text:
        header_parts.append(f'<p class="reader__subtitle">{html_mod.escape(text)}</p>')
    elif text.startswith('Documento Estratégico'):
        header_parts.append(f'<p class="reader__meta">{html_mod.escape(text)}</p>')
    elif text.startswith('Autor:'):
        header_parts.append(f'<p class="reader__meta">{html_mod.escape(text)}</p>')
    elif text.startswith('Classificação:'):
        header_parts.append(f'<p class="reader__meta">{html_mod.escape(text)}</p>')
    else:
        header_parts.append(render_elem(e))
header_html = '\n'.join(header_parts)

# ── Render TOC ──
toc_html = '<h2>Índice</h2>\n<ol class="toc">\n'
for entry in toc_entries:
    indent = '  ' if entry['level'] == 'h3' else ''
    toc_html += f"{indent}<li><a href=\"#{entry['anchor']}\">{html_mod.escape(entry['text'])}</a></li>\n"
toc_html += '</ol>'

# ── Final HTML ──
full_html = f'''<!DOCTYPE html>
<html lang="pt-PT">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Portugal no Mundo Multipolar — UNIR</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet">
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{
  font-family:'Inter',system-ui,-apple-system,sans-serif;
  font-size:16px;line-height:1.75;color:#212121;
  max-width:780px;margin:0 auto;padding:48px 28px;background:#fff
}}
h2{{
  font-family:'Inter',system-ui,sans-serif;
  font-size:1.5rem;font-weight:700;color:#0D47A1;
  margin:44px 0 12px;padding-bottom:8px;
  border-bottom:2px solid #E3F2FD
}}
.reader__title{{
  font-size:2rem;font-weight:800;color:#0D47A1;
  border-bottom:3px solid #0D47A1;padding-bottom:12px;margin-bottom:8px
}}
h3{{
  font-family:'Inter',system-ui,sans-serif;
  font-size:1.15rem;font-weight:600;color:#1565C0;
  margin:32px 0 10px
}}
h4{{
  font-family:'Inter',system-ui,sans-serif;
  font-size:1.05rem;font-weight:600;color:#333;
  margin:24px 0 8px
}}
p{{margin:0 0 14px}}
ul,ol{{padding-left:22px;margin:0 0 18px}}
li{{margin-bottom:5px}}
strong{{color:#0D47A1;font-weight:600}}
a{{color:#0D47A1;text-decoration:none}}
a:hover{{text-decoration:underline}}
table{{
  width:100%;border-collapse:collapse;margin:20px 0;font-size:0.9rem
}}
th,td{{
  border:1px solid #B0BEC5;padding:10px 14px;
  text-align:left;vertical-align:top
}}
th{{
  background:#0D47A1;color:#fff;font-weight:600;
  font-family:'Inter',sans-serif;font-size:0.85rem
}}
tr:nth-child(even){{background:#F5F7FA}}
blockquote{{
  border-left:4px solid #0D47A1;margin:22px 0;padding:12px 20px;
  background:#E3F2FD;color:#333;border-radius:0 8px 8px 0
}}
hr{{border:none;border-top:1px solid #B0BEC5;margin:32px 0}}
.reader__org{{
  text-transform:uppercase;letter-spacing:3px;
  font-size:0.8rem;font-weight:700;color:#9E9E9E;margin-bottom:6px
}}
.reader__subtitle{{
  font-size:1.15rem;color:#1565C0;font-weight:500;margin-bottom:4px
}}
.reader__meta{{
  font-size:0.85rem;color:#757575;margin-bottom:1px
}}
.toc{{
  background:#F5F7FA;border-radius:10px;padding:20px 28px;margin:20px 0 32px
}}
.toc li{{margin-bottom:4px;font-size:0.95rem}}
.toc li li{{margin-left:18px;font-size:0.9rem}}
@media(max-width:640px){{
  body{{padding:28px 14px;font-size:15px}}
  h2{{font-size:1.25rem}}
  .reader__title{{font-size:1.5rem}}
  h3{{font-size:1.05rem}}
  table{{font-size:0.8rem}}
  th,td{{padding:7px 10px}}
}}
</style>
</head>
<body>

{header_html}

{toc_html}

{body_html}

</body>
</html>'''

with open(OUT, 'w', encoding='utf-8') as f:
    f.write(full_html)

print(f"Generated: {OUT}")
print(f"Size: {len(full_html)} bytes")
print(f"H2: {full_html.count('<h2 ')}  H3: {full_html.count('<h3 ')}  H4: {full_html.count('<h4 ')}")
print(f"Tables: {full_html.count('<table>')}")
