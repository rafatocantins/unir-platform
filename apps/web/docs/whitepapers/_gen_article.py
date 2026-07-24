#!/usr/bin/env python3
"""Generate reader-style HTML article + shared CSS from .docx."""
from docx import Document
import re, html as html_mod

DOC = '/root/unir-platform/apps/web/docs/whitepapers/WHITEPAPER_PORTUGAL_MUNDO_MULTIPOLAR_UNIR.docx'
OUT_HTML = '/root/unir-platform/apps/web/docs/whitepapers/WHITEPAPER_PORTUGAL_MUNDO_MULTIPOLAR_UNIR.html'
OUT_CSS = '/root/unir-platform/apps/web/css/reader.css'

doc = Document(DOC)

import unicodedata

def slugify(text):
    t = text.lower()
    t = unicodedata.normalize('NFKD', t).encode('ASCII', 'ignore').decode()
    t = re.sub(r'[^a-z0-9\s-]', '', t)
    t = re.sub(r'\s+', '-', t.strip())
    return t

# ── Parse document into structured sections ──
raw = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name if p.style else 'None'
    raw.append({'idx': i, 'style': style, 'text': text})

sections = []
stack = [{'level': 0, 'children': sections}]
heading_counts = {}
current_content = []

def flush_content(target):
    if current_content:
        merged = []
        i = 0
        while i < len(current_content):
            item = current_content[i]
            if item['type'] == 'bullet':
                bullets = []
                while i < len(current_content) and current_content[i]['type'] == 'bullet':
                    bullets.append(current_content[i]['text'])
                    i += 1
                merged.append({'type': 'ul', 'items': bullets})
            else:
                merged.append(item)
                i += 1
        target['content_before_children'] = merged
        current_content.clear()

for r in raw:
    text = r['text']
    style = r['style']
    if not text:
        continue

    if style.startswith('Heading'):
        level = int(style.split()[-1])
        anchor = slugify(text)
        if anchor in heading_counts:
            heading_counts[anchor] += 1
            anchor = f"{anchor}-{heading_counts[anchor]}"
        else:
            heading_counts[anchor] = 0

        sec = {'title': text, 'anchor': anchor, 'level': level,
               'children': [], 'content_before_children': []}

        while stack[-1]['level'] >= level:
            flush_content(stack[-1])
            stack.pop()

        parent = stack[-1]
        parent['children'].append(sec)
        stack.append(sec)
        current_content = sec['content_before_children']

    elif style in ('List Bullet', 'List Bullet 2'):
        current_content.append({'type': 'bullet', 'text': text})
    else:
        current_content.append({'type': 'p', 'text': text})

while len(stack) > 1:
    flush_content(stack[-1])
    stack.pop()
flush_content(stack[0])

# ── Extract header, body sections ──
root_content = stack[0].get('content_before_children', [])
root_children = sections

body_sections = []
toc_section = None

for sec in root_children:
    if sec['title'] in ('Indice', 'Índice'):
        toc_section = sec
    elif sec['level'] == 1 and sec['title'] not in ('UNIR', 'PORTUGAL NO MUNDO MULTIPOLAR'):
        body_sections.append(sec)

# ── Tables ──
tables_data = []
for ti, table in enumerate(doc.tables):
    rows = []
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        rows.append(cells)
    tables_data.append(rows)

TABLE_PLACEMENT = {
    0: '2-contexto-geopolitico-o-fim-do-unipolarismo',
    1: 'dependencias-atuais',
    2: 'oportunidades-inexploradas',
    3: '51-oportunidades',
    4: '82-vantagem-competitiva-portuguesa',
    5: 'fase-1-2026-2028---fundacao',
    6: 'fase-2-2029-2032---expansao',
    7: 'fase-3-2033-2035---consolidacao',
    8: '11-analise-de-riscos-e-mitigacao',
}

# ── Render helpers ──
def render_table(rows):
    if not rows:
        return ''
    html_rows = ''
    for ri, row in enumerate(rows):
        tag = 'th' if ri == 0 else 'td'
        cells = ''.join(f'<{tag}>{html_mod.escape(c)}</{tag}>' for c in row)
        html_rows += f'<tr>{cells}</tr>\n'
    return f'<table class="reader__table">\n{html_rows}</table>'

def render_content(items, section_anchor=''):
    lines = []
    for item in items:
        if item['type'] == 'p':
            lines.append(f'<p>{html_mod.escape(item["text"])}</p>')
        elif item['type'] == 'ul':
            lis = '\n'.join(f'<li>{html_mod.escape(it)}</li>' for it in item['items'])
            lines.append(f'<ul>\n{lis}\n</ul>')

    for ti, anchor in TABLE_PLACEMENT.items():
        if anchor == section_anchor:
            lines.append(render_table(tables_data[ti]))

    return '\n'.join(lines)

def render_section(sec, is_toplevel=True):
    anchor = sec['anchor']
    title = html_mod.escape(sec['title'])
    level = sec['level']
    content_html = render_content(sec.get('content_before_children', []), anchor)

    children_html = ''
    for child in sec.get('children', []):
        children_html += render_section(child, is_toplevel=False)

    if is_toplevel and level == 1:
        return f'''<details class="reader__section" open id="{anchor}">
<summary class="reader__section-title">{title}</summary>
<div class="reader__section-body">
{content_html}
{children_html}
</div>
</details>'''
    elif level == 2:
        return f'<h4 class="reader__subsection" id="{anchor}">{title}</h4>\n{content_html}\n{children_html}'
    elif level == 3:
        return f'<h5 class="reader__subsubsection" id="{anchor}">{title}</h5>\n{content_html}\n{children_html}'
    else:
        return f'{content_html}\n{children_html}'

# ── Header ──
header_html_parts = []
for item in root_content:
    text = item.get('text', '') if item['type'] == 'p' else ''
    if not text:
        continue
    if text == 'UNIR':
        header_html_parts.append(f'<p class="reader__org">{html_mod.escape(text)}</p>')
    elif 'Estratégia Lusófona' in text:
        header_html_parts.append(f'<p class="reader__subtitle">{html_mod.escape(text)}</p>')
    elif text.startswith('Documento Estratégico'):
        header_html_parts.append(f'<p class="reader__meta">{html_mod.escape(text)}</p>')
    elif text.startswith('Autor:'):
        header_html_parts.append(f'<p class="reader__meta">{html_mod.escape(text)}</p>')
    elif text.startswith('Classificação:'):
        header_html_parts.append(f'<p class="reader__meta">{html_mod.escape(text)}</p>')

title_html = '<h1 class="reader__title">PORTUGAL NO MUNDO MULTIPOLAR</h1>'
header_html = '\n'.join(header_html_parts)

# ── TOC for sidebar ──
toc_items = []
for sec in body_sections:
    if sec['level'] == 1:
        toc_items.append({'title': sec['title'], 'anchor': sec['anchor'], 'level': 1})
        for child in sec.get('children', []):
            if child['level'] == 2:
                toc_items.append({'title': child['title'], 'anchor': child['anchor'], 'level': 2})

sidebar_toc = ''
for item in toc_items:
    cls = 'reader__toc-link--sub' if item['level'] == 2 else ''
    sidebar_toc += f'<li class="{cls}"><a href="#{item["anchor"]}" class="reader__toc-link">{html_mod.escape(item["title"])}</a></li>\n'

# ── Body ──
body_html_parts = []
for sec in body_sections:
    body_html_parts.append(render_section(sec, is_toplevel=True))
body_html = '\n'.join(body_html_parts)

# ── Article HTML ──
article_html = f'''<div class="reader__progress" id="progressBar"><div class="reader__progress-fill" id="progressFill"></div></div>

<article class="reader">

<aside class="reader__sidebar">
  <nav class="reader__toc">
    <h4 class="reader__toc-title">Indice</h4>
    <ol>
{sidebar_toc}
    </ol>
  </nav>
</aside>

<div class="reader__body">

<header class="reader__header">
{header_html}
{title_html}
</header>

<div class="reader__content">
{body_html}
</div>

<footer class="reader__footer">
  <p>Documento Estrategico &middot; UNIR &mdash; Unidos pela Nacao, Inovacao e Responsabilidade</p>
  <p>Maio 2026 &middot; Classificacao: Publico</p>
</footer>

</div>
</article>

<script>
(function(){{
  var tocLinks = document.querySelectorAll('.reader__toc-link');
  var headings = document.querySelectorAll('.reader__section[id], .reader__subsection[id]');
  if (!headings.length) return;
  var observer = new IntersectionObserver(function(entries){{
    entries.forEach(function(e){{
      if (e.isIntersecting) {{
        tocLinks.forEach(function(l){{ l.classList.remove('reader__toc-link--active'); }});
        var link = document.querySelector('.reader__toc-link[href="#' + e.target.id + '"]');
        if (link) link.classList.add('reader__toc-link--active');
      }}
    }});
  }}, {{root: document.getElementById('readerModal'), rootMargin: '-10% 0px -75% 0px'}});
  headings.forEach(function(h){{ observer.observe(h); }});
}})();
</script>'''

with open(OUT_HTML, 'w', encoding='utf-8') as f:
    f.write(article_html)

# ── Shared CSS ──
css_content = '''/* ══════════════════════════════════════════════════════
   UNIR Reader CSS — Shared for all article/blog documents
   ══════════════════════════════════════════════════════ */

.reader__progress {
  position: fixed; top: 0; left: 0; right: 0; height: 4px;
  background: #E0E0E0; z-index: 2000;
}
.reader__progress-fill {
  height: 100%; width: 0;
  background: linear-gradient(90deg, #0D47A1, #1565C0);
  transition: width 0.15s linear;
}

.reader {
  display: grid;
  grid-template-columns: 240px 1fr;
  gap: 48px;
  max-width: 1100px;
  margin: 0 auto;
  padding: 56px 32px 80px;
  font-family: 'Inter', system-ui, -apple-system, sans-serif;
  color: #212121;
  line-height: 1.75;
}

.reader__sidebar {
  position: sticky; top: 80px; align-self: start;
  max-height: calc(100vh - 120px); overflow-y: auto;
}
.reader__toc {
  background: #F5F7FA;
  border-radius: 12px;
  padding: 20px 18px;
}
.reader__toc-title {
  font-size: 0.8rem; font-weight: 700; text-transform: uppercase;
  letter-spacing: 1.5px; color: #9E9E9E;
  margin-bottom: 12px;
}
.reader__toc ol {
  list-style: none; padding: 0; margin: 0;
}
.reader__toc li {
  margin-bottom: 2px;
}
.reader__toc li.reader__toc-link--sub {
  padding-left: 14px;
}
.reader__toc-link {
  display: block; padding: 4px 8px; border-radius: 6px;
  font-size: 0.85rem; color: #616161; text-decoration: none;
  transition: all 0.15s;
}
.reader__toc-link:hover {
  background: #E3F2FD; color: #0D47A1;
}
.reader__toc-link--active {
  background: #0D47A1; color: #fff !important;
  font-weight: 600;
}

.reader__body {
  min-width: 0;
}

.reader__header {
  margin-bottom: 44px; padding-bottom: 32px;
  border-bottom: 1px solid #E0E0E0;
}
.reader__org {
  text-transform: uppercase; letter-spacing: 3px;
  font-size: 0.8rem; font-weight: 700; color: #9E9E9E;
  margin-bottom: 8px;
}
.reader__title {
  font-size: 2.2rem; font-weight: 900; color: #0D47A1;
  line-height: 1.2; margin: 0 0 8px;
}
.reader__subtitle {
  font-size: 1.15rem; color: #1565C0; font-weight: 500;
  margin-bottom: 4px;
}
.reader__meta {
  font-size: 0.85rem; color: #757575; margin-bottom: 2px;
}

.reader__section {
  margin-bottom: 16px;
  border: 1px solid #E0E0E0;
  border-radius: 12px;
  overflow: hidden;
  background: #fff;
  transition: box-shadow 0.2s;
}
.reader__section[open] {
  box-shadow: 0 2px 12px rgba(0,0,0,0.06);
}
.reader__section-title {
  display: flex; align-items: center; gap: 10px;
  padding: 16px 22px;
  font-size: 1.15rem; font-weight: 700; color: #0D47A1;
  cursor: pointer; user-select: none;
  background: linear-gradient(135deg, #F5F7FA, #fff);
  list-style: none;
  transition: background 0.15s;
}
.reader__section-title::-webkit-details-marker { display: none; }
.reader__section-title::before {
  content: '+'; display: inline-flex; align-items: center; justify-content: center;
  width: 22px; height: 22px; border-radius: 50%;
  background: #E3F2FD; color: #0D47A1;
  font-weight: 700; font-size: 1rem; flex-shrink: 0;
  transition: transform 0.2s;
}
.reader__section[open] .reader__section-title::before {
  content: '-'; background: #0D47A1; color: #fff;
}
.reader__section-title:hover {
  background: #E3F2FD;
}
.reader__section-body {
  padding: 4px 22px 22px;
}

.reader__subsection {
  font-size: 1.05rem; font-weight: 700; color: #1565C0;
  margin: 28px 0 10px; padding-top: 8px;
}
.reader__subsubsection {
  font-size: 0.95rem; font-weight: 600; color: #333;
  margin: 20px 0 6px;
}

.reader__body p {
  margin: 0 0 14px; font-size: 1rem;
}
.reader__body ul, .reader__body ol {
  padding-left: 22px; margin: 0 0 18px;
}
.reader__body li {
  margin-bottom: 5px;
}
.reader__body strong {
  color: #0D47A1; font-weight: 600;
}

.reader__metric-grid {
  display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
  gap: 14px; margin: 22px 0;
}
.reader__metric-card {
  background: linear-gradient(135deg, #E3F2FD, #F5F7FA);
  border-radius: 10px; padding: 18px;
  border-left: 4px solid #0D47A1;
}
.reader__metric-card-icon {
  font-size: 1.5rem; margin-bottom: 6px;
}
.reader__metric-card-title {
  font-size: 0.85rem; font-weight: 700; color: #0D47A1;
  margin-bottom: 4px;
}
.reader__metric-card-text {
  font-size: 0.85rem; color: #424242; line-height: 1.4;
}

.reader__table {
  width: 100%; border-collapse: collapse;
  margin: 22px 0; font-size: 0.88rem;
  border-radius: 8px; overflow: hidden;
  box-shadow: 0 1px 6px rgba(0,0,0,0.06);
}
.reader__table th {
  background: #0D47A1; color: #fff; font-weight: 600;
  padding: 10px 14px; text-align: left; font-size: 0.82rem;
}
.reader__table td {
  border-bottom: 1px solid #E0E0E0;
  padding: 10px 14px; text-align: left; vertical-align: top;
}
.reader__table tr:hover td {
  background: #E3F2FD;
}
.reader__table tr:nth-child(even) td {
  background: #FAFAFA;
}
.reader__table tr:nth-child(even):hover td {
  background: #E3F2FD;
}

.reader__sidenote {
  float: right; width: 180px; margin: 6px -24px 14px 20px;
  padding: 12px; background: #FFF8E1; border-radius: 8px;
  font-size: 0.8rem; color: #5D4037; line-height: 1.4;
  border-left: 3px solid #FF8F00;
}
.reader__sidenote strong {
  display: block; font-size: 1.3rem; color: #0D47A1; margin-bottom: 2px;
}

.reader__footer {
  margin-top: 56px; padding-top: 24px;
  border-top: 1px solid #E0E0E0;
  text-align: center;
}
.reader__footer p {
  font-size: 0.82rem; color: #9E9E9E; margin-bottom: 4px;
}

@media (max-width: 900px) {
  .reader {
    grid-template-columns: 1fr;
    gap: 24px;
    padding: 32px 16px 60px;
  }
  .reader__sidebar {
    position: static; max-height: none; order: -1;
  }
  .reader__toc {
    display: flex; flex-wrap: wrap; gap: 6px;
    padding: 12px 14px;
  }
  .reader__toc-title { width: 100%; margin-bottom: 4px; }
  .reader__toc ol { display: flex; flex-wrap: wrap; gap: 4px; }
  .reader__toc li { margin: 0; }
  .reader__toc li.reader__toc-link--sub { padding-left: 0; }
  .reader__toc-link { font-size: 0.78rem; padding: 3px 8px; }
  .reader__title { font-size: 1.5rem; }
  .reader__section-title { font-size: 1rem; padding: 12px 16px; }
  .reader__section-body { padding: 4px 16px 16px; }
}
'''

with open(OUT_CSS, 'w', encoding='utf-8') as f:
    f.write(css_content)

print(f"OK: {OUT_HTML} ({len(article_html)} bytes)")
print(f"OK: {OUT_CSS} ({len(css_content)} bytes)")
print(f"Sections: {len(body_sections)}")
print(f"H2: {article_html.count('<summary class=')}, H4: {article_html.count('<h4 class=')}")
print(f"Tables: {article_html.count('<table class=')}")
